from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_DOCX = Path("output/docx/robbinsdale-room-pack-list.docx")


ROOMS = [
    {
        "name": "Living Room",
        "pack_sections": [
            (
                "Furniture",
                [
                    "Cream sofa",
                    "(2) Brown swivel chairs",
                    "Round wood coffee table",
                    "Fluted wood console table",
                    "8 x 10 neutral rug",
                    "Black floor lamp",
                ],
            ),
            (
                "Textiles",
                [
                    "Sofa: (2) 24 x 24 natural linen pillows",
                    "Sofa: (2) 20 x 20 Loloi Irene Ivory pillows",
                    "Sofa: (1) neutral textured lumbar",
                    "Sofa: neutral throw over one arm",
                    "Each chair: (1) 20 x 20 Loloi Irene Ivory pillow",
                ],
            ),
            (
                "Decor Pull",
                [
                    "Coffee table: wood tray, 2 books, real candle, small greenery",
                    "Console table: black vase, greenery, books, small object",
                    "Mantel: large vase, greenery, books, small object",
                    "Top shelf: leaning landscape artwork, small plant",
                    "Bottom shelf: books, decorative bowl, small object",
                ],
            ),
        ],
        "placement": [
            "Center the rug in the room.",
            "Center the coffee table on the rug.",
            "Center the sofa directly across from the fireplace.",
            "Center the console table behind the sofa.",
            "Place one swivel chair on each side of the fireplace.",
            "Angle both chairs slightly toward the coffee table.",
            "Place the floor lamp in the window corner.",
        ],
        "styling_sections": [
            (
                "Styling Notes",
                [
                    "Keep the shelves minimal with open space.",
                    "Maintain symmetry around the fireplace.",
                    "Keep tabletop styling light and neutral.",
                ],
            )
        ],
    },
    {
        "name": "Dining Room",
        "pack_sections": [
            (
                "Furniture",
                [
                    "Dining table",
                    "(4) Cream upholstered chairs",
                ],
            ),
            (
                "Wall Decor",
                [
                    "36 x 36 framed artwork for the back wall",
                    '36 to 40 in round mirror for the wall between the windows',
                ],
            ),
            (
                "Table Styling",
                [
                    "Choose one only: wood bowl",
                    "Or choose one only: ceramic vase with greenery",
                ],
            ),
        ],
        "placement": [
            "Center the table under the chandelier.",
            "Use four chairs only.",
            "Center the artwork on the back wall.",
            "Center the mirror between the windows.",
        ],
        "styling_sections": [
            (
                "Styling Notes",
                [
                    "Nothing else goes on the table.",
                    "Keep the room simple and balanced.",
                ],
            )
        ],
    },
    {
        "name": "Primary Bedroom",
        "pack_sections": [
            (
                "Furniture",
                [
                    "Queen upholstered bed",
                    "(2) Nightstands",
                    "(2) Lamps",
                    "8 x 10 rug",
                ],
            ),
            (
                "Bedding",
                [
                    "White duvet",
                    "White sheets",
                    "(2) Euro pillows with neutral shams",
                    "(2) Standard sleeping pillows",
                    "(1) Neutral lumbar",
                    "Knit throw folded across the lower third of the bed",
                ],
            ),
            (
                "Decor Pull",
                [
                    "Artwork for over the bed",
                    "Large olive tree for corner placement",
                ],
            ),
        ],
        "placement": [
            "Center the bed on the main wall.",
            "Keep the nightstands even on both sides.",
            "Center one lamp on each nightstand.",
            "Center the artwork over the bed.",
            "Place the olive tree in the open corner.",
        ],
        "styling_sections": [
            (
                "Styling Notes",
                [
                    "No dresser in this room.",
                    "Keep the bed styling crisp and hotel-clean.",
                ],
            )
        ],
    },
    {
        "name": "Porch",
        "pack_sections": [
            (
                "Furniture",
                [
                    "Black metal loveseat",
                    "(2) Black metal chairs",
                    "Black metal coffee table",
                    "Porch rug",
                ],
            ),
            (
                "Cushions + Pillows",
                [
                    "Cream seat cushions",
                    "Cream back cushions",
                    "Loveseat: (2) ivory textured pillows",
                    "Loveseat: (1) neutral lumbar",
                    "Each chair: (1) sage pillow",
                ],
            ),
            (
                "Decor Pull",
                [
                    "Coffee table: tray, real candle, small faux plant",
                    "Hanging ferns",
                    "Large planter",
                    "Lantern",
                    "Welcome mat",
                ],
            ),
        ],
        "placement": [
            "Center the rug with the seating group.",
            "Anchor the loveseat as the main piece.",
            "Flank the loveseat with the two chairs.",
            "Center the coffee table within the seating area.",
            "Use the planter and lantern to finish the entry side.",
        ],
        "styling_sections": [
            (
                "Styling Notes",
                [
                    "Keep the porch feeling clean and welcoming.",
                    "Use greenery to soften the black metal pieces.",
                ],
            )
        ],
    },
]


ACCENT = RGBColor(86, 72, 60)
INK = RGBColor(31, 31, 31)
MUTED = RGBColor(103, 103, 103)
LINE = "D8D3CD"
FILL = "F4F1ED"


def set_run_font(run, name, size, bold=False, color=INK, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rfonts = run._element.rPr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        run._element.rPr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_spacing(paragraph, before=0, after=0, line=260):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line / 240


def add_text_paragraph(container, text, size=9.5, bold=False, color=INK, italic=False, before=0, after=1):
    paragraph = container.add_paragraph()
    set_paragraph_spacing(paragraph, before=before, after=after, line=255)
    run = paragraph.add_run(text)
    set_run_font(run, "Aptos", size, bold=bold, color=color, italic=italic)
    return paragraph


def add_section_label(container, text):
    paragraph = container.add_paragraph()
    set_paragraph_spacing(paragraph, before=0, after=4, line=240)
    shade_paragraph(paragraph, FILL)
    run = paragraph.add_run(f"  {text.upper()}  ")
    set_run_font(run, "Aptos", 8.5, bold=True, color=ACCENT)
    return paragraph


def add_room_page(doc, room, index, total):
    if index:
        doc.add_page_break()

    header = doc.add_paragraph()
    set_paragraph_spacing(header, before=0, after=2, line=240)
    run = header.add_run(f"AJ HOME | ROBBINSDALE | ROOM {index + 1} OF {total}")
    set_run_font(run, "Aptos", 8.5, bold=True, color=MUTED)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=2, line=240)
    run = title.add_run(room["name"])
    set_run_font(run, "Aptos Display", 20, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, before=0, after=8, line=240)
    run = subtitle.add_run("Warm Transitional | Less Is More | Symmetrical")
    set_run_font(run, "Aptos", 9, italic=True, color=MUTED)

    divider = doc.add_paragraph()
    set_paragraph_spacing(divider, before=0, after=8, line=240)
    run = divider.add_run(" ")
    p_pr = divider._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), LINE)
    p_bdr.append(bottom)

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    remove_table_borders(table)

    left, right = table.rows[0].cells
    left.width = Inches(3.15)
    right.width = Inches(3.15)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(left)
    set_cell_margins(right)

    left.paragraphs[0]._element.getparent().remove(left.paragraphs[0]._element)
    right.paragraphs[0]._element.getparent().remove(right.paragraphs[0]._element)

    add_section_label(left, "Pack List")
    for label, items in room["pack_sections"]:
        add_text_paragraph(left, label, size=9.5, bold=True, color=ACCENT, before=1, after=1)
        for item in items:
            add_text_paragraph(left, item, size=9.25, after=1)

    add_section_label(right, "Placement")
    for item in room["placement"]:
        add_text_paragraph(right, item, size=9.25, after=1)

    add_section_label(right, "Styling + Notes")
    for label, items in room["styling_sections"]:
        add_text_paragraph(right, label, size=9.5, bold=True, color=ACCENT, before=1, after=1)
        for item in items:
            add_text_paragraph(right, item, size=9.25, after=1)


def configure_document(doc):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")


def build_document():
    doc = Document()
    configure_document(doc)
    for index, room in enumerate(ROOMS):
        add_room_page(doc, room, index, len(ROOMS))
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
